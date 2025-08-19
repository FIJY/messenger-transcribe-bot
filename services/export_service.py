# services/export_service.py - Полная система экспорта файлов
import logging
import os
import tempfile
import zipfile
from datetime import datetime
from typing import Dict, Any, Optional, List
import asyncio

# Импорты для разных форматов
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfutils
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase import pdfmetrics

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class ExportService:
    """Сервис для экспорта транскрипций в разные форматы"""

    def __init__(self):
        self.temp_dir = "/tmp/exports"
        os.makedirs(self.temp_dir, exist_ok=True)
        logger.info("🗂️ ExportService инициализирован")

    async def export_transcription(self, transcription_data: Dict[str, Any],
                                   format_type: str, include_metadata: bool = True) -> Optional[str]:
        """
        Экспорт транскрипции в указанном формате

        Args:
            transcription_data: Данные транскрипции
            format_type: Тип формата (txt, docx, pdf, srt)
            include_metadata: Включать ли метаданные

        Returns:
            str: Путь к созданному файлу или None при ошибке
        """
        try:
            text = transcription_data.get('text', '')
            language = transcription_data.get('language', 'unknown')

            if not text.strip():
                logger.error("Пустой текст для экспорта")
                return None

            # Генерируем уникальное имя файла
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename_base = f"transcription_{timestamp}"

            if format_type == 'txt':
                return await self._export_txt(text, language, filename_base, include_metadata, transcription_data)
            elif format_type == 'docx':
                return await self._export_docx(text, language, filename_base, include_metadata, transcription_data)
            elif format_type == 'pdf':
                return await self._export_pdf(text, language, filename_base, include_metadata, transcription_data)
            elif format_type == 'srt':
                return await self._export_srt(text, filename_base, transcription_data)
            else:
                logger.error(f"Неподдерживаемый формат: {format_type}")
                return None

        except Exception as e:
            logger.error(f"Ошибка экспорта в формат {format_type}: {e}", exc_info=True)
            return None

    async def _export_txt(self, text: str, language: str, filename_base: str,
                          include_metadata: bool, transcription_data: Dict[str, Any]) -> str:
        """Экспорт в текстовый файл"""
        filepath = os.path.join(self.temp_dir, f"{filename_base}.txt")

        content = []

        if include_metadata:
            content.extend([
                "📝 ТРАНСКРИПЦИЯ",
                "=" * 50,
                f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
                f"Язык: {language.upper()}",
                f"Длина текста: {len(text)} символов",
                ""
            ])

            # Добавляем статистику
            if language in ['zh', 'ja', 'ko']:
                content.append(f"Количество символов: {len(text)}")
            else:
                word_count = len(text.split())
                content.append(f"Количество слов: {word_count}")

            content.extend(["", "ТЕКСТ:", "-" * 20, ""])

        content.append(text)

        # Записываем файл
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))

        logger.info(f"✅ TXT файл создан: {filepath}")
        return filepath

    async def _export_docx(self, text: str, language: str, filename_base: str,
                           include_metadata: bool, transcription_data: Dict[str, Any]) -> Optional[str]:
        """Экспорт в Word документ"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx не установлен")
            return None

        try:
            filepath = os.path.join(self.temp_dir, f"{filename_base}.docx")

            # Создаем документ
            doc = Document()

            if include_metadata:
                # Заголовок
                heading = doc.add_heading('📝 Транскрипция', 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Метаданные в таблице
                table = doc.add_table(rows=4, cols=2)
                table.style = 'Table Grid'

                table.cell(0, 0).text = 'Дата создания'
                table.cell(0, 1).text = datetime.now().strftime('%d.%m.%Y %H:%M')

                table.cell(1, 0).text = 'Язык'
                table.cell(1, 1).text = language.upper()

                table.cell(2, 0).text = 'Длина текста'
                table.cell(2, 1).text = f"{len(text)} символов"

                if language in ['zh', 'ja', 'ko']:
                    table.cell(3, 0).text = 'Количество символов'
                    table.cell(3, 1).text = str(len(text))
                else:
                    word_count = len(text.split())
                    table.cell(3, 0).text = 'Количество слов'
                    table.cell(3, 1).text = str(word_count)

                # Разделитель
                doc.add_paragraph()
                doc.add_heading('Текст транскрипции', level=1)

            # Основной текст
            # Разбиваем на параграфы для лучшего форматирования
            paragraphs = text.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
                else:
                    doc.add_paragraph()  # Пустая строка

            # Сохраняем
            doc.save(filepath)

            logger.info(f"✅ DOCX файл создан: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"Ошибка создания DOCX: {e}")
            return None

    async def _export_pdf(self, text: str, language: str, filename_base: str,
                          include_metadata: bool, transcription_data: Dict[str, Any]) -> Optional[str]:
        """Экспорт в PDF файл"""
        if not PDF_AVAILABLE:
            logger.error("reportlab не установлен")
            return None

        try:
            filepath = os.path.join(self.temp_dir, f"{filename_base}.pdf")

            # Создаем PDF документ
            doc = SimpleDocTemplate(filepath, pagesize=A4,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=18)

            # Стили
            styles = getSampleStyleSheet()

            # Кастомные стили
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Центр
            )

            content = []

            if include_metadata:
                # Заголовок
                content.append(Paragraph("📝 Транскрипция", title_style))
                content.append(Spacer(1, 12))

                # Метаданные
                metadata_style = styles['Normal']
                content.extend([
                    Paragraph(f"<b>Дата создания:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}", metadata_style),
                    Paragraph(f"<b>Язык:</b> {language.upper()}", metadata_style),
                    Paragraph(f"<b>Длина текста:</b> {len(text)} символов", metadata_style),
                ])

                if language in ['zh', 'ja', 'ko']:
                    content.append(Paragraph(f"<b>Количество символов:</b> {len(text)}", metadata_style))
                else:
                    word_count = len(text.split())
                    content.append(Paragraph(f"<b>Количество слов:</b> {word_count}", metadata_style))

                content.append(Spacer(1, 20))
                content.append(Paragraph("<b>Текст транскрипции:</b>", styles['Heading2']))
                content.append(Spacer(1, 12))

            # Основной текст
            text_style = styles['Normal']
            paragraphs = text.split('\n')

            for para in paragraphs:
                if para.strip():
                    # Экранируем HTML символы для reportlab
                    safe_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    content.append(Paragraph(safe_para, text_style))
                    content.append(Spacer(1, 6))

            # Строим PDF
            doc.build(content)

            logger.info(f"✅ PDF файл создан: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"Ошибка создания PDF: {e}")
            return None

    async def _export_srt(self, text: str, filename_base: str,
                          transcription_data: Dict[str, Any]) -> str:
        """Экспорт в формат субтитров SRT"""
        filepath = os.path.join(self.temp_dir, f"{filename_base}.srt")

        # Получаем длительность файла
        duration_seconds = transcription_data.get('duration_seconds', 0)

        # Разбиваем текст на предложения для субтитров
        sentences = self._split_text_for_subtitles(text)

        if not sentences:
            sentences = [text]  # Fallback

        # Рассчитываем время для каждого субтитра
        if duration_seconds > 0:
            time_per_subtitle = duration_seconds / len(sentences)
        else:
            time_per_subtitle = 3  # 3 секунды по умолчанию

        srt_content = []

        for i, sentence in enumerate(sentences):
            start_time = i * time_per_subtitle
            end_time = (i + 1) * time_per_subtitle

            # Форматируем время в SRT формат (HH:MM:SS,mmm)
            start_srt = self._seconds_to_srt_time(start_time)
            end_srt = self._seconds_to_srt_time(end_time)

            srt_content.extend([
                str(i + 1),  # Номер субтитра
                f"{start_srt} --> {end_srt}",  # Временной интервал
                sentence.strip(),  # Текст
                ""  # Пустая строка
            ])

        # Записываем файл
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(srt_content))

        logger.info(f"✅ SRT файл создан: {filepath}")
        return filepath

    def _split_text_for_subtitles(self, text: str, max_chars: int = 60) -> List[str]:
        """Разбивает текст на подходящие для субтитров части"""
        sentences = []

        # Разбиваем по предложениям
        import re
        sentence_endings = re.split(r'[.!?]+', text)

        for sentence in sentence_endings:
            sentence = sentence.strip()
            if not sentence:
                continue

            # Если предложение слишком длинное, разбиваем по запятым
            if len(sentence) > max_chars:
                parts = sentence.split(',')
                current_part = ""

                for part in parts:
                    if len(current_part + part) <= max_chars:
                        current_part += part + ","
                    else:
                        if current_part:
                            sentences.append(current_part.rstrip(','))
                        current_part = part + ","

                if current_part:
                    sentences.append(current_part.rstrip(','))
            else:
                sentences.append(sentence)

        return [s for s in sentences if s.strip()]

    def _seconds_to_srt_time(self, seconds: float) -> str:
        """Конвертирует секунды в SRT временной формат"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds - int(seconds)) * 1000)

        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"

    async def create_zip_archive(self, file_paths: List[str], archive_name: str) -> Optional[str]:
        """Создает ZIP архив из списка файлов"""
        try:
            archive_path = os.path.join(self.temp_dir, f"{archive_name}.zip")

            with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in file_paths:
                    if os.path.exists(file_path):
                        # Добавляем только имя файла, без пути
                        filename = os.path.basename(file_path)
                        zipf.write(file_path, filename)
                        logger.info(f"Добавлен в архив: {filename}")
                    else:
                        logger.warning(f"Файл не найден: {file_path}")

            logger.info(f"✅ ZIP архив создан: {archive_path}")
            return archive_path

        except Exception as e:
            logger.error(f"Ошибка создания ZIP архива: {e}")
            return None

    async def export_all_formats(self, transcription_data: Dict[str, Any]) -> Optional[str]:
        """Экспорт во всех доступных форматах и создание ZIP архива"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            exported_files = []

            # Экспортируем в каждом формате
            formats = ['txt', 'srt']  # Всегда доступные форматы

            if DOCX_AVAILABLE:
                formats.append('docx')
            if PDF_AVAILABLE:
                formats.append('pdf')

            for format_type in formats:
                file_path = await self.export_transcription(transcription_data, format_type)
                if file_path and os.path.exists(file_path):
                    exported_files.append(file_path)

            if not exported_files:
                logger.error("Не удалось создать ни одного файла")
                return None

            # Создаем ZIP архив
            archive_path = await self.create_zip_archive(exported_files, f"transcription_all_{timestamp}")

            # Очищаем временные файлы
            for file_path in exported_files:
                try:
                    os.remove(file_path)
                except Exception as e:
                    logger.warning(f"Не удалось удалить временный файл {file_path}: {e}")

            return archive_path

        except Exception as e:
            logger.error(f"Ошибка при экспорте всех форматов: {e}")
            return None

    async def cleanup_old_files(self, max_age_hours: int = 24):
        """Очистка старых экспортированных файлов"""
        try:
            current_time = datetime.now().timestamp()
            cleaned_count = 0

            for filename in os.listdir(self.temp_dir):
                file_path = os.path.join(self.temp_dir, filename)

                if os.path.isfile(file_path):
                    file_age_hours = (current_time - os.path.getctime(file_path)) / 3600

                    if file_age_hours > max_age_hours:
                        os.remove(file_path)
                        cleaned_count += 1
                        logger.debug(f"Удален старый файл: {filename}")

            if cleaned_count > 0:
                logger.info(f"🧹 Очищено {cleaned_count} старых файлов экспорта")

        except Exception as e:
            logger.error(f"Ошибка очистки старых файлов: {e}")

    def get_available_formats(self) -> Dict[str, bool]:
        """Возвращает список доступных форматов экспорта"""
        return {
            'txt': True,  # Всегда доступен
            'srt': True,  # Всегда доступен
            'docx': DOCX_AVAILABLE,
            'pdf': PDF_AVAILABLE
        }

    def get_format_description(self, format_type: str) -> str:
        """Возвращает описание формата"""
        descriptions = {
            'txt': 'Обычный текстовый файл с метаданными',
            'docx': 'Документ Microsoft Word с форматированием',
            'pdf': 'PDF отчет с красивым оформлением',
            'srt': 'Файл субтитров для видео'
        }
        return descriptions.get(format_type, 'Неизвестный формат')